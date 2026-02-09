"""
Document Extractor Service - Extracts text content from PDF, Word, Excel files.

This service processes uploaded documents locally to extract structured text,
saving tokens by not sending binary data to the LLM.
"""

import io
from pathlib import Path
from typing import Optional, Tuple
from dataclasses import dataclass


@dataclass
class ExtractionResult:
    """Result of document extraction."""
    success: bool
    text: str
    document_type: str
    filename: str
    error: Optional[str] = None
    metadata: Optional[dict] = None


class DocumentExtractor:
    """Extracts text content from various document formats."""

    SUPPORTED_EXTENSIONS = {
        '.pdf': 'PDF',
        '.docx': 'Word',
        '.doc': 'Word',
        '.xlsx': 'Excel',
        '.xls': 'Excel',
        '.txt': 'Text',
        '.csv': 'CSV',
    }

    @classmethod
    def detect_document_type(cls, filename: str, content_type: Optional[str] = None) -> Optional[str]:
        """Detect document type from filename or content type."""
        extension = Path(filename).suffix.lower()
        return cls.SUPPORTED_EXTENSIONS.get(extension)

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Check if the file type is supported for extraction."""
        extension = Path(filename).suffix.lower()
        return extension in cls.SUPPORTED_EXTENSIONS

    @staticmethod
    def extract_pdf(content: bytes, filename: str) -> ExtractionResult:
        """Extract text from PDF file."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(content))
            text_parts = []

            for i, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"--- Page {i} ---\n{page_text.strip()}")

            full_text = "\n\n".join(text_parts)

            metadata = {}
            if reader.metadata:
                metadata = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "pages": len(reader.pages)
                }

            return ExtractionResult(
                success=True,
                text=full_text,
                document_type="PDF",
                filename=filename,
                metadata=metadata
            )
        except Exception as e:
            return ExtractionResult(
                success=False,
                text="",
                document_type="PDF",
                filename=filename,
                error=str(e)
            )

    @staticmethod
    def extract_word(content: bytes, filename: str) -> ExtractionResult:
        """Extract text from Word document."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(content))
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_rows.append(" | ".join(row_text))
                if table_rows:
                    text_parts.append("\n[Tableau]\n" + "\n".join(table_rows))

            full_text = "\n".join(text_parts)

            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }

            return ExtractionResult(
                success=True,
                text=full_text,
                document_type="Word",
                filename=filename,
                metadata=metadata
            )
        except Exception as e:
            return ExtractionResult(
                success=False,
                text="",
                document_type="Word",
                filename=filename,
                error=str(e)
            )

    @staticmethod
    def extract_excel(content: bytes, filename: str) -> ExtractionResult:
        """Extract text from Excel file."""
        try:
            from openpyxl import load_workbook

            workbook = load_workbook(io.BytesIO(content), data_only=True)
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_content = [f"--- Feuille: {sheet_name} ---"]

                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    # Only keep rows with at least one non-empty value
                    if any(val.strip() for val in row_values):
                        sheet_content.append(" | ".join(row_values))

                if len(sheet_content) > 1:  # More than just the header
                    text_parts.append("\n".join(sheet_content))

            full_text = "\n\n".join(text_parts)

            metadata = {
                "sheets": workbook.sheetnames,
                "sheet_count": len(workbook.sheetnames)
            }

            return ExtractionResult(
                success=True,
                text=full_text,
                document_type="Excel",
                filename=filename,
                metadata=metadata
            )
        except Exception as e:
            return ExtractionResult(
                success=False,
                text="",
                document_type="Excel",
                filename=filename,
                error=str(e)
            )

    @staticmethod
    def extract_text(content: bytes, filename: str) -> ExtractionResult:
        """Extract text from plain text file."""
        text = None
        # Try multiple encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            text = content.decode('utf-8', errors='ignore')

        return ExtractionResult(
            success=True,
            text=text,
            document_type="Text",
            filename=filename
        )

    @classmethod
    def extract(cls, content: bytes, filename: str, content_type: Optional[str] = None) -> ExtractionResult:
        """
        Extract text from a document based on its type.

        Args:
            content: Binary content of the file
            filename: Name of the file (used for type detection)
            content_type: Optional MIME type

        Returns:
            ExtractionResult with extracted text or error
        """
        extension = Path(filename).suffix.lower()

        if extension == '.pdf':
            return cls.extract_pdf(content, filename)
        elif extension in ['.docx', '.doc']:
            return cls.extract_word(content, filename)
        elif extension in ['.xlsx', '.xls']:
            return cls.extract_excel(content, filename)
        elif extension in ['.txt', '.csv']:
            return cls.extract_text(content, filename)
        else:
            # Try to decode as text
            try:
                text = content.decode('utf-8')
                return ExtractionResult(
                    success=True,
                    text=text,
                    document_type="Text",
                    filename=filename
                )
            except:
                return ExtractionResult(
                    success=False,
                    text="",
                    document_type="Unknown",
                    filename=filename,
                    error=f"Type de fichier non supporté: {extension}"
                )
